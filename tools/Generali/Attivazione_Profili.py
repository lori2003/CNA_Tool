import os
import shutil
import json
import time
from datetime import datetime
from pathlib import Path
import streamlit as st
import re

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import pythoncom
    import comtypes.client
    _COM_AVAILABLE = True
except ImportError:
    _COM_AVAILABLE = False


# -----------------------------------------------------------------------------
# CONFIGURAZIONE PATH E COSTANTI
# -----------------------------------------------------------------------------
BASE_TOOL_DIR = Path(__file__).resolve().parent / "Modulo - Attivazione Profili"
TEMPLATE_PATH = BASE_TOOL_DIR / "-Lettera di designazione e di istruzioni GDPR per collaboratori territorio.docx"
DIR_IN_SOSPESO = BASE_TOOL_DIR / "In Sospeso"
CONFIG_FILE = BASE_TOOL_DIR / "config_percorsi.json"

# Percorso di default degli archiviati
DEFAULT_ARCHIVE_PATH = r"F:\Cna Pensionati\CNA PENSIONATI 2026\RICHIESTE CREDENZIALI OPERATORI"

# Assicura esistenza cartelle base
DIR_IN_SOSPESO.mkdir(parents=True, exist_ok=True)

def load_config():
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f).get("archive_path", DEFAULT_ARCHIVE_PATH)
        except:
            return DEFAULT_ARCHIVE_PATH
    return DEFAULT_ARCHIVE_PATH

def save_config(new_path):
    with open(CONFIG_FILE, "w") as f:
        json.dump({"archive_path": new_path}, f)

# -----------------------------------------------------------------------------
# TOOL DEFINITION
# -----------------------------------------------------------------------------
TOOL = {
    'id': 'attivazione_profili',
    'name': 'Attivazione Profili (PDF)',
    'description': (
        '#### 📌 1. FINALITÀ DEL TOOL\n'
        "Automatizza il workflow di 'attivazione' e regolarizzazione dei collaboratori territorio. Gestisce "
        "sia la creazione di documenti ufficiali (Lettere GDPR) tramite template, sia l'importazione massiva "
        'di file già esistenti, organizzandoli in una struttura di cartelle standardizzata.\n\n'
        '#### 🚀 2. COME UTILIZZARLO\n'
        '* **Crea nuovo profilo:** Inserisci Nome, Cognome e Tipo Cartella → il tool genera il PDF dalla lettera GDPR.\n'
        '* **Importa file esistenti:** Carica i file (PDF, Word, ecc.) → il tool crea le cartelle leggendo i nomi dai filename.\n\n'
        '#### 📂 4. RISULTATO FINALE\n'
        'Cartelle strutturate in `In Sospeso/` pronte per la validazione o l\'archiviazione finale.'
    ),
    'version': '1.6.0',
    'params': [
        {
            'key': 'modalita',
            'label': 'Modalità',
            'type': 'radio',
            'options': ['Crea nuovo profilo', 'Importa file esistenti'],
            'default': 'Crea nuovo profilo',
        },
        {
            'key': 'nome',
            'label': 'Nome',
            'type': 'text',
            'placeholder': 'Es: Mario',
        },
        {
            'key': 'cognome',
            'label': 'Cognome',
            'type': 'text',
            'placeholder': 'Es: Rossi',
        },
        {
            'key': 'tipo_cartella',
            'label': 'Tipo Cartella',
            'type': 'radio',
            'options': ['Bancadati', 'INPS', 'BD e INPS'],
            'default': 'Bancadati',
        },
    ],
    'inputs': [
        {
            'key': 'files',
            'label': 'File da importare *(solo modalità "Importa file esistenti")*',
            'type': 'file_multi',
            'required': False,
        },
    ],
}

# -----------------------------------------------------------------------------
# LOGICA DI BUSINESS
# -----------------------------------------------------------------------------
def genera_nuovo_profilo(nome, cognome, tipo_cartella):
    """Crea cartella, genera PDF e garantisce l'eliminazione del Word."""
    if not _COM_AVAILABLE or not _DOCX_AVAILABLE:
        st.error(
            "❌ Questo tool richiede **Microsoft Word** e le librerie COM di Windows. "
            "Non è compatibile con Linux/Cloud — usalo solo sul tuo PC locale."
        )
        return False

    if not TEMPLATE_PATH.exists():
        st.error(f"Template non trovato in: {TEMPLATE_PATH}")
        return False

    folder_name = f"{nome.strip()} {cognome.strip()} ({tipo_cartella})"
    dest_folder = DIR_IN_SOSPESO / folder_name
    
    if dest_folder.exists():
        ts = datetime.now().strftime("%H%M%S")
        dest_folder = DIR_IN_SOSPESO / f"{folder_name}_{ts}"
    
    try:
        dest_folder.mkdir(parents=True, exist_ok=True)
        docx_temp = dest_folder / f"temp_elaborazione_{int(time.time())}.docx"
        pdf_dest = dest_folder / TEMPLATE_PATH.with_suffix(".pdf").name
        
        # 1. Copia e Modifica Word
        shutil.copy2(TEMPLATE_PATH, docx_temp)
        doc = Document(str(docx_temp))
        
        for p in doc.paragraphs:
            if "Gent. Sig.r/ ra" in p.text:
                p.text = p.text.replace("Gent. Sig.r/ ra", f"Gent. Sig.r/ ra {nome} {cognome},")
                break
        
        # Sostituzione data
        conteggio = 0
        data_oggi = datetime.now().strftime("%d/%m/%Y")
        for p in doc.paragraphs:
            if "Roma," in p.text:
                conteggio += 1
                if conteggio == 2:
                    p.text = p.text.replace("Roma,", f"Roma, {data_oggi}")
                    break
        doc.save(str(docx_temp))

        # 2. Conversione PDF con pulizia COM
        pythoncom.CoInitialize()
        try:
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            abs_docx = str(docx_temp.resolve())
            abs_pdf = str(pdf_dest.resolve())
            
            doc_word = word.Documents.Open(abs_docx)
            doc_word.SaveAs(abs_pdf, FileFormat=17)
            doc_word.Close(0) # 0 = wdDoNotSaveChanges
            word.Quit()
            
            # Rilascio oggetti per permettere eliminazione
            del doc_word
            del word
            time.sleep(1.0) # Attesa per rilascio file system
        finally:
            pythoncom.CoUninitialize()

            
        # 3. ELIMINAZIONE CERTIFICATA DEL WORD
        if docx_temp.exists():
            try:
                os.remove(docx_temp)
            except:
                # Tentativo finale ritardato se bloccato
                time.sleep(2.0)
                if docx_temp.exists(): 
                    try: os.remove(docx_temp)
                    except: pass
            
        return True
    except Exception as e:
        st.error(f"Errore tecnico: {e}")
        return False

def importa_profilo_esistente(nome, cognome, tipo_cartella, uploaded_files):
    """Crea la cartella e salva i file caricati senza generare il template."""
    folder_name = f"{nome.strip()} {cognome.strip()} ({tipo_cartella})"
    dest_folder = DIR_IN_SOSPESO / folder_name
    
    if dest_folder.exists():
        ts = datetime.now().strftime("%H%M%S")
        dest_folder = DIR_IN_SOSPESO / f"{folder_name}_{ts}"
    
    try:
        # Crea la cartella solo se non esiste già (evita doppioni da esecuzioni multiple)
        if not dest_folder.exists():
            dest_folder.mkdir(parents=True, exist_ok=True)
            time.sleep(0.5) # Piccolo delay per permettere al file system di stabilizzarsi
            
        for f in uploaded_files:
            file_path = dest_folder / f.name
            if not file_path.exists():
                with open(file_path, "wb") as out_f:
                    out_f.write(f.getvalue())
        return True
    except Exception as e:
        st.error(f"Errore durante l'importazione: {e}")
        return False

def estrai_nome_cognome_da_filename(filename):
    """Tenta di indovinare nome e cognome dal nome del file."""
    # Rimuove estensione e pulisce caratteri speciali
    base = os.path.splitext(filename)[0]
    # Sostituisce trattini, underscore e punti con spazi
    base = base.replace("_", " ").replace("-", " ").replace(".", " ")
    # Rimuove numeri (es. date o codici)
    base = re.sub(r'\d+', ' ', base)
    # Divide per spazi e ripulisce
    parti = [p.strip().capitalize() for p in base.split() if len(p.strip()) > 1]
    
    nome, cognome = "", ""
    if len(parti) >= 2:
        # Se abbiamo almeno due parti, prendiamo le prime due come nome e cognome
        nome, cognome = parti[0], parti[1]
    elif len(parti) == 1:
        nome = parti[0]
    return nome, cognome

# -----------------------------------------------------------------------------
# INTERFACCIA UTENTE (DASHBOARD)
# -----------------------------------------------------------------------------
def get_ui_top():
    # --- INIZIALIZZAZIONE SESSION STATE ---
    if "import_profiles" not in st.session_state:
        st.session_state["import_profiles"] = {}
    if "last_import_sig" not in st.session_state:
        st.session_state["last_import_sig"] = ""
    if "import_in_progress" not in st.session_state:
        st.session_state["import_in_progress"] = False

    # CSS per nascondere pulsante standard e gestire i bordi rossi
    st.markdown("""
        <style>
        /* Nasconde il pulsante Esegui di default */
        div[data-testid="stButton"] button[key^="run_"] { display: none !important; }
        
        /* Forza il bordo rosso sui container con border=True */
        div[data-testid="stVerticalBlockBordered"] {
            border: 2px solid #ff4b4b !important;
            border-radius: 12px !important;
        }
        
        /* Stile per i blocchi file */
        .file-block {
            background-color: rgba(255, 75, 75, 0.05);
            padding: 10px;
            border-radius: 8px;
            margin-bottom: 15px;
            border-left: 4px solid #ff4b4b;
        }
        </style>
    """, unsafe_allow_html=True)

    # --- SETTINGS ARCHIVIO (Fisicamente dentro il bordo) ---
    archive_path = load_config()
    with st.container(border=True):
        st.markdown("<h3 style='text-align: center; margin-top: -10px;'>📂 DESTINAZIONE ARCHIVIO</h3>", unsafe_allow_html=True)
        c_path, c_save, c_open = st.columns([0.6, 0.2, 0.2])
        with c_path:
            new_path = st.text_input("Path", value=archive_path, label_visibility="collapsed", key="cfg_path")
        with c_save:
            if st.button("💾 Salva", use_container_width=True, key="save_cfg_btn"):
                save_config(new_path)
                st.rerun()
        with c_open:
            if st.button("📂 Vai", use_container_width=True, key="open_cfg_btn"):
                os.system(f'explorer "{new_path}"')

    st.write("") # Spaciatore

    # --- SEZIONE CREAZIONE ---
    st.markdown("## ✨ Crea Nuovo Profilo")
    with st.container(border=True):
        col1, col2 = st.columns(2)
        with col1:
            n = st.text_input("Nome", placeholder="Es: Mario", key="manual_nome")
        with col2:
            c = st.text_input("Cognome", placeholder="Es: Rossi", key="manual_cognome")
        t = st.radio("Tipo Cartella", ["Bancadati", "INPS", "BD e INPS"], horizontal=True)
        if st.button("Crea Cartella", type="primary", use_container_width=True, key="main_gen_btn"):
            if n and c:
                with st.spinner("Generazione PDF e pulizia in corso..."):
                    if genera_nuovo_profilo(n, c, t):
                        st.success(f"PDF creato per {n} {c}")
                        st.rerun()
            else:
                st.warning("⚠️ Inserisci sia Nome che Cognome.")

    # --- SEZIONE IMPORTAZIONE ---
    st.markdown("## 📤 Importa File Già Pronti")
    with st.container(border=True):
        # Sposto il caricatore in alto per un drop immediato
        uploaded_files = st.file_uploader("1. Trascina qui i file (PDF, Word, etc.)", 
                                         accept_multiple_files=True, 
                                         key="upload_import")
        
        if uploaded_files:
            # --- LOGICA DI RILEVAMENTO AUTOMATICO E GESTIONE STATO ---
            files_sig = "|".join([f"{f.name}_{f.size}" for f in uploaded_files])
            
            # Se i file sono cambiati, aggiorniamo la lista profili senza perdere i dati esistenti
            if st.session_state["last_import_sig"] != files_sig:
                current_profiles = st.session_state["import_profiles"]
                new_profiles = {}
                
                for f in uploaded_files:
                    if f.name in current_profiles:
                        # Manteniamo quello che l'utente ha già scritto
                        new_profiles[f.name] = current_profiles[f.name]
                    else:
                        # Nuovo file: rileviamo nome e cognome
                        n_auto, c_auto = estrai_nome_cognome_da_filename(f.name)
                        new_profiles[f.name] = {
                            "nome": n_auto,
                            "cognome": c_auto,
                            "tipo": "Bancadati"
                        }
                
                st.session_state["import_profiles"] = new_profiles
                st.session_state["last_import_sig"] = files_sig
                st.rerun()

            st.info(f"✅ {len(uploaded_files)} file pronti per l'importazione.")
            
            # --- IMPOSTAZIONI GLOBALI ---
            with st.expander("⚙️ Impostazioni Globali (Applica a tutti)", expanded=False):
                col_glob_opt, col_glob_val = st.columns([0.4, 0.6])
                with col_glob_opt:
                    apply_all = st.checkbox("Applica lo stesso tipo a tutti", value=False, key="apply_all_toggle")
                with col_glob_val:
                    global_type = st.radio("Tipo Cartella Globale", ["Bancadati", "INPS", "BD e INPS"], 
                                         horizontal=True, key="global_type", disabled=not apply_all)

            st.write("---")
            
            # --- LOOP SUI FILE ---
            for i, f in enumerate(uploaded_files):
                fname = f.name
                if fname not in st.session_state["import_profiles"]:
                    continue # Sicurezza
                
                profile = st.session_state["import_profiles"][fname]
                
                # Container per singolo file
                st.markdown(f"<div class='file-block'>📄 <b>File:</b> {fname}</div>", unsafe_allow_html=True)
                
                col_n, col_swap, col_c, col_t = st.columns([0.28, 0.08, 0.28, 0.36])
                
                # SPOSTIAMO LO SWAP PRIMA DEGLI INPUT NEL CODICE
                with col_swap:
                    st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)
                    if st.button("🔄", key=f"swap_{fname}", help="Scambia Nome e Cognome", use_container_width=True):
                        # Prendiamo i valori attuali dai widget (o dal profilo se non ancora toccati)
                        n_current = st.session_state.get(f"n_{fname}", profile["nome"])
                        c_current = st.session_state.get(f"c_{fname}", profile["cognome"])
                        
                        # Invertiamo nel dizionario principale
                        st.session_state["import_profiles"][fname]["nome"] = c_current
                        st.session_state["import_profiles"][fname]["cognome"] = n_current
                        
                        # Aggiorniamo le chiavi widget PRIMA del rerun per sicurezza
                        st.session_state[f"n_{fname}"] = c_current
                        st.session_state[f"c_{fname}"] = n_current
                        st.rerun()

                with col_n:
                    new_n = st.text_input(f"Nome", value=profile["nome"], key=f"n_{fname}")
                    st.session_state["import_profiles"][fname]["nome"] = new_n
                
                with col_c:
                    new_c = st.text_input(f"Cognome", value=profile["cognome"], key=f"c_{fname}")
                    st.session_state["import_profiles"][fname]["cognome"] = new_c
                
                with col_t:
                    current_type = global_type if apply_all else profile["tipo"]
                    new_t = st.radio(f"Tipo", ["Bancadati", "INPS", "BD e INPS"], 
                                   index=["Bancadati", "INPS", "BD e INPS"].index(current_type),
                                   horizontal=True, key=f"t_{fname}", 
                                   disabled=apply_all,
                                   label_visibility="visible")
                    if not apply_all:
                        st.session_state["import_profiles"][fname]["tipo"] = new_t

            st.write("")
            if st.button(f"Crea {len(uploaded_files)} Cartelle e Importa Tutto", type="primary", use_container_width=True, key="import_all_btn"):
                if not st.session_state.get("import_in_progress", False):
                    st.session_state["import_in_progress"] = True
                    successos = 0
                    errors = []
                    
                    with st.spinner("Creazione profili in corso..."):
                        for f in uploaded_files:
                            p = st.session_state["import_profiles"][f.name]
                            p_type = global_type if apply_all else p["tipo"]
                            
                            if p["nome"] and p["cognome"]:
                                if importa_profilo_esistente(p["nome"], p["cognome"], p_type, [f]):
                                    successos += 1
                                else:
                                    errors.append(f.name)
                            else:
                                errors.append(f"{f.name} (Dati mancanti)")
                    
                    if successos > 0:
                        st.success(f"✅ Importati con successo {successos} profili.")
                        # RESET TOTALE PER EVITARE DOPPIONI
                        st.session_state["import_profiles"] = {}
                        st.session_state["last_import_sig"] = "RESET_DONE"
                        st.session_state["import_in_progress"] = False
                        time.sleep(1.5)
                        st.rerun()
                    
                    st.session_state["import_in_progress"] = False
        else:
            st.caption("Trascina i file sopra per iniziare l'importazione multipla.")

    st.divider()

    # --- SEZIONE DASHBOARD ---
    st.markdown("## 📂 Profili In Sospeso")
    search_query = st.text_input("🔎 Ricerca nominativo...", placeholder="Cerca tra i sospesi...").lower()

    if not DIR_IN_SOSPESO.exists():
        st.info("Nessuna cartella trovata.")
        return

    all_items = sorted([p for p in DIR_IN_SOSPESO.iterdir() if p.is_dir()], key=lambda p: p.stat().st_ctime, reverse=True)
    filtered_items = [p for p in all_items if search_query in p.name.lower()] if search_query else all_items

    if not filtered_items:
        if search_query:
            st.warning(f"Nessun risultato per: '{search_query}'")
        else:
            st.info("📭 Nessun operatore in sospeso.")
    else:
        for folder in filtered_items:
            with st.container(border=True):
                col_info, col_open, col_arch, col_del = st.columns([0.46, 0.18, 0.18, 0.18])
                with col_info:
                    st.markdown(f"**{folder.name}**")
                    ctime = datetime.fromtimestamp(folder.stat().st_ctime).strftime("%d/%m/%Y %H:%M")
                    st.caption(f"📅 {ctime}")
                with col_open:
                    if st.button("📂 Apri", key=f"open_{folder.name}", use_container_width=True):
                        os.system(f'explorer "{folder}"') 
                with col_arch:
                    if st.button("🗄️ Archiva", key=f"arch_{folder.name}", use_container_width=True):
                        try:
                            target_archive = load_config()
                            dest_base = Path(target_archive)
                            dest_base.mkdir(parents=True, exist_ok=True)
                            dest = dest_base / folder.name
                            if dest.exists(): dest = dest_base / f"{folder.name}_{datetime.now().strftime('%H%M%S')}"
                            shutil.move(str(folder), str(dest))
                            st.rerun()
                        except Exception as e:
                            st.error(f"Errore archivio: {e}")
                with col_del:
                    if st.button("🗑️", key=f"del_{folder.name}", use_container_width=True, help="Elimina definitivamente"):
                        try:
                            shutil.rmtree(folder)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Errore eliminazione: {e}")

def run(out_dir, modalita="Crea nuovo profilo", nome="", cognome="",
        tipo_cartella="Bancadati", files=None, **kwargs):
    if modalita == "Crea nuovo profilo":
        if not nome or not cognome:
            raise ValueError("Nome e Cognome sono obbligatori per la creazione del profilo.")

        success = genera_nuovo_profilo(nome.strip(), cognome.strip(), tipo_cartella)
        if not success:
            raise RuntimeError(
                f"Errore nella generazione del profilo per {nome} {cognome}. "
                "Verifica che Microsoft Word sia installato e che il template esista."
            )

        prefix = f"{nome.strip()} {cognome.strip()} ({tipo_cartella})"
        candidates = sorted(
            [d for d in DIR_IN_SOSPESO.iterdir() if d.is_dir() and d.name.startswith(prefix)],
            key=lambda d: d.stat().st_ctime, reverse=True,
        )
        if not candidates:
            raise RuntimeError("Profilo generato ma cartella non trovata in 'In Sospeso'.")
        return [f for f in candidates[0].rglob("*") if f.is_file()]

    else:  # Importa file esistenti
        if files is None:
            raise ValueError("Nessun file caricato per l'importazione.")
        if not isinstance(files, list):
            files = [files]

        output_files = []
        for file_path in files:
            if not isinstance(file_path, Path):
                continue

            if nome and cognome:
                n, c = nome.strip(), cognome.strip()
            else:
                n, c = estrai_nome_cognome_da_filename(file_path.name)

            if not n or not c:
                continue

            folder_name = f"{n} {c} ({tipo_cartella})"
            dest_folder = DIR_IN_SOSPESO / folder_name
            if dest_folder.exists():
                ts = datetime.now().strftime("%H%M%S")
                dest_folder = DIR_IN_SOSPESO / f"{folder_name}_{ts}"
            dest_folder.mkdir(parents=True, exist_ok=True)

            dest_file = dest_folder / file_path.name
            shutil.copy2(file_path, dest_file)
            output_files.append(dest_file)

        if not output_files:
            raise ValueError(
                "Nessun file importato. Assicurati di aver caricato i file "
                "e che i nomi contengano almeno Nome e Cognome."
            )
        return output_files
